"""Round trip DOCUMENTACION.md through md2docx and check nothing got lost.

Expectations are counted from the markdown itself, not hardcoded, so the test
survives edits to the document and only fails if the converter drops something.

    .venv/bin/python docs/test_md2docx.py       # or: pytest docs/test_md2docx.py
"""

import os
import re
import sys
import tempfile

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from md2docx import MONO, convert  # noqa: E402

from docx import Document  # noqa: E402

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
MD   = os.path.join(ROOT, "DOCUMENTACION.md")


def _plain(text):
    """Markdown text as it should read once converted: no links, no emphasis markers."""
    text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)
    return re.sub(r"\*\*(.+?)\*\*|\*(.+?)\*|`(.+?)`|~~(.+?)~~",
                  lambda m: next(g for g in m.groups() if g is not None), text)


def _md_blocks(lines):
    """(headings, image_paths, table_count, code_block_count) as written in the markdown."""
    headings, images, tables, code = [], [], 0, 0
    in_code = False
    prev_table = False
    for raw in lines:
        s = raw.strip()
        if s.startswith("```"):
            in_code = not in_code
            if in_code:
                code += 1
            continue
        if in_code:
            continue
        m = re.match(r"^(#{1,6})\s+(.*)", s)
        if m:
            headings.append((min(len(m.group(1)), 3), m.group(2)))
        m = re.fullmatch(r"!\[([^\]]*)\]\(<?([^>)]+?)>?\)", s)
        if m:
            images.append(m.group(2))
        is_table = s.startswith("|") and s.endswith("|")
        if is_table and not prev_table:
            tables += 1
        prev_table = is_table
    return headings, images, tables, code


def test_documentacion():
    with open(MD, encoding="utf-8") as fh:
        lines = fh.read().splitlines()
    headings, images, n_tables, n_code = _md_blocks(lines)

    for img in images:
        assert os.path.exists(os.path.join(ROOT, img)), f"imagen ausente: {img}"

    with tempfile.TemporaryDirectory() as tmp:
        doc = Document(convert(MD, os.path.join(tmp, "out.docx")))

    texts  = [p.text for p in doc.paragraphs]
    styled = [(p.style.name, p.text) for p in doc.paragraphs]

    # headings: same titles, same levels, same order
    got = [(int(s[-1]), t) for s, t in styled if s.startswith("Heading ")]
    want = [(lvl, _plain(t)) for lvl, t in headings]
    assert got == want, f"headings\n got={got}\nwant={want}"

    assert len(doc.inline_shapes) == len(images), "faltan imágenes en el docx"
    assert len(doc.tables) == n_tables, f"tablas: {len(doc.tables)} != {n_tables}"

    # every table: bold grey header row that repeats, no ragged rows
    for t in doc.tables:
        hdr = t.rows[0]
        assert hdr.cells[0].paragraphs[0].runs[0].bold
        assert hdr._tr.trPr.findall(
            "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tblHeader"
        )
        for row in t.rows:
            assert len(row.cells) == len(t.columns)

    # the numbered index survived with its numbers, and links lost their URLs
    assert any(t.startswith("1.  Qué hace") for t in texts)
    all_text = texts + [c.text for t in doc.tables for r in t.rows for c in r.cells]
    assert not any("](" in t for t in all_text), "quedó sintaxis de enlace sin convertir"

    # blockquote and code blocks
    assert any(p.style.name == "Normal" and p.runs and p.runs[0].italic
               for p in doc.paragraphs), "blockquote en cursiva"
    mono = [p for p in doc.paragraphs
            if p.runs and p.runs[0].font.name == MONO and p.runs[0].font.size.pt == 8.5]
    assert len(mono) >= n_code, f"líneas de código: {len(mono)} para {n_code} bloques"

    # inline formatting reached the body
    assert any(r.bold for p in doc.paragraphs for r in p.runs)


def test_bullets():
    """DOCUMENTACION.md ya no usa viñetas; el conversor sí las soporta."""
    with tempfile.TemporaryDirectory() as tmp:
        md = os.path.join(tmp, "b.md")
        with open(md, "w", encoding="utf-8") as fh:
            fh.write("- uno\n  - anidada\n")
        texts = [p.text for p in Document(convert(md)).paragraphs]
    assert texts[0].startswith("•") and texts[1].startswith("○")


if __name__ == "__main__":
    test_documentacion()
    test_bullets()
    print("ok")

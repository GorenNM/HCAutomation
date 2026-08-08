#!/usr/bin/env python3
"""Convert Markdown reports to .docx using MMN_Template.docx as style reference.

Usage:
    python docs/_generate_docx.py                        # converts DEFAULT_REPORTS
    python docs/_generate_docx.py path/to/file.md        # single file
    python docs/_generate_docx.py inventario_datasets    # named target
    python docs/_generate_docx.py guia_stewards          # named target

Output: <same_directory>/<same_stem>.docx
"""

from __future__ import annotations

import os
import re
import sys
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# ── Paths ──────────────────────────────────────────────────────────────────────
# Script lives in docs/_toolkit/; docs/ is the parent.
_SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
_DOCS_DIR   = os.path.dirname(_SCRIPT_DIR)          # docs/
TEMPLATE    = os.path.join(_SCRIPT_DIR, "MMN_Template.docx")

_WHITE = RGBColor(0xFF, 0xFF, 0xFF)
_BLACK = RGBColor(0x00, 0x00, 0x00)

# ── Inline formatting ──────────────────────────────────────────────────────────
_INLINE = re.compile(
    r"\*\*(.+?)\*\*"
    r"|\*(.+?)\*"
    r"|`(.+?)`"
    r"|~~(.+?)~~"
)
_TABLE_SEP = re.compile(r"^\|[-|\s:]+\|$")
_IMAGE_RE  = re.compile(r"!\[([^\]]*)\]\(<?([^>)]+?)>?\)")
_LINK_RE   = re.compile(r"\[([^\]]+)\]\([^)]+\)")
_KV_RE     = re.compile(r"^\*\*([^*]+?):\*\*\s*(.*)")


def _add_runs(para, text: str) -> None:
    pos = 0
    for m in _INLINE.finditer(text):
        if m.start() > pos:
            para.add_run(text[pos:m.start()])
        bold, italic, code, strike = m.group(1), m.group(2), m.group(3), m.group(4)
        if bold:
            r = para.add_run(bold);   r.bold = True
        elif italic:
            r = para.add_run(italic); r.italic = True
        elif code:
            r = para.add_run(code);   r.font.name = "Courier New"; r.font.size = Pt(9)
        elif strike:
            r = para.add_run(strike); r.font.strike = True
        pos = m.end()
    if pos < len(text):
        para.add_run(text[pos:])


def _strip_links(text: str) -> str:
    return _LINK_RE.sub(r"\1", text)


def _cell_bg(cell, hex_color: str) -> None:
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def _cell_borders(cell) -> None:
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    bdr  = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"),   "single")
        el.set(qn("w:sz"),    "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "000000")
        bdr.append(el)
    tcPr.append(bdr)


def _set_table_borders(tbl) -> None:
    """Set full borders on the table including inside vertical and horizontal lines."""
    tbl_el = tbl._tbl
    tblPr  = tbl_el.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl_el.insert(0, tblPr)
    tblBorders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"),   "single")
        el.set(qn("w:sz"),    "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "000000")
        tblBorders.append(el)
    tblPr.append(tblBorders)


def _repeat_header(row) -> None:
    """Mark a table row to repeat as header on every page."""
    trPr = row._tr.get_or_add_trPr()
    hdr  = OxmlElement("w:tblHeader")
    trPr.append(hdr)


def _add_table(doc, headers: list[str], rows: list[list[str]]) -> None:
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = "Normal Table"
    _set_table_borders(tbl)
    # Header row — #DCDCDC background, black bold text, repeat on every page
    hdr_row = tbl.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        _cell_borders(cell)
        _cell_bg(cell, "DCDCDC")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(h.strip())
        r.bold = True; r.font.color.rgb = _BLACK; r.font.size = Pt(9)
    _repeat_header(hdr_row)
    # Data rows — no background, full borders
    for row in rows:
        tr = tbl.add_row()
        for i, val in enumerate(row):
            cell = tr.cells[i]
            _cell_borders(cell)
            p = cell.paragraphs[0]
            _add_runs(p, _strip_links(val.strip()))
            for r in p.runs:
                r.font.size = Pt(9)
    doc.add_paragraph()


def _add_code_block(doc, code_lines: list[str]) -> None:
    for line in code_lines:
        p   = doc.add_paragraph(style="Normal")
        p.paragraph_format.left_indent  = Inches(0.3)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        run = p.add_run(line if line else " ")
        run.font.name = "Courier New"; run.font.size = Pt(8.5)
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"),   "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"),  "F5F5F5")
        pPr.append(shd)
    doc.add_paragraph()


def _clear_body(doc) -> None:
    body    = doc.element.body
    sect_pr = body.find(qn("w:sectPr"))
    for child in list(body):
        if child is not sect_pr:
            body.remove(child)


def _fix_styles(doc) -> None:
    """Normalise heading and body styles to Metro reference: black, Calibri, correct sizes."""
    FONT = "Calibri"
    overrides = {
        "Heading 1":      (18, True),
        "Heading 2":      (13, True),
        "Heading 3":      (11, True),
        "Normal":         (10, False),
        "List Paragraph": (10, False),
    }
    for name, (size, bold) in overrides.items():
        try:
            s = doc.styles[name]
            s.font.name       = FONT
            s.font.size       = Pt(size)
            s.font.color.rgb  = _BLACK
            s.font.bold       = bold if bold else None
        except KeyError:
            pass


def _add_cover(doc, title: str, kv: dict[str, str]) -> None:
    """Render cover page: title → Requerimiento → Versión → Fecha → Control de versiones."""
    # Title
    p = doc.add_paragraph(style="Heading 1")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(title)

    # Requerimiento
    rq = kv.get("Requerimiento", "")
    if rq:
        lbl = doc.add_paragraph(style="Normal")
        lbl.alignment = WD_ALIGN_PARAGRAPH.CENTER
        lbl.add_run("Requerimiento:").font.size = Pt(10)

        val = doc.add_paragraph(style="Normal")
        val.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = val.add_run(rq)
        r.bold = True; r.font.size = Pt(12)

    # Versión
    ver = kv.get("Versión", kv.get("Version", ""))
    if ver:
        p_ver = doc.add_paragraph(style="Normal")
        p_ver.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p_ver.add_run(f"Versión {ver}")
        r.bold = True; r.font.size = Pt(11)

    # Fecha
    fecha = kv.get("Fecha", "")
    if fecha:
        p_f = doc.add_paragraph(style="Normal")
        p_f.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_f.add_run(fecha).font.size = Pt(10)

    doc.add_paragraph()  # spacer

    # Control de versiones table
    p_cv = doc.add_paragraph(style="Normal")
    r = p_cv.add_run("Control de versiones")
    r.bold = True; r.font.size = Pt(11)

    tbl = doc.add_table(rows=1, cols=4)
    tbl.style = "Normal Table"
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    _set_table_borders(tbl)
    hdr_row = tbl.rows[0]
    for i, h in enumerate(["Fecha", "Versión", "Descripción", "Autor"]):
        cell = hdr_row.cells[i]
        _cell_borders(cell)
        _cell_bg(cell, "DCDCDC")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(h)
        r.bold = True; r.font.color.rgb = _BLACK; r.font.size = Pt(9)
    _repeat_header(hdr_row)

    # First row pre-populated
    row1 = tbl.add_row()
    for i, val in enumerate([fecha, ver, "Creación del documento", kv.get("Autores", "")]):
        _cell_borders(row1.cells[i])
        row1.cells[i].paragraphs[0].add_run(val).font.size = Pt(9)

    # Two blank rows
    for _ in range(2):
        br = tbl.add_row()
        for cell in br.cells:
            _cell_borders(cell)
            cell.paragraphs[0].add_run("").font.size = Pt(9)

    # Page break after cover
    doc.add_page_break()


# ── Main converter ─────────────────────────────────────────────────────────────

def convert(md_path: str, output_path: str | None = None, template: str = TEMPLATE) -> str:
    with open(md_path, encoding="utf-8") as fh:
        lines = fh.read().splitlines()

    doc = Document(template)
    _clear_body(doc)
    _fix_styles(doc)

    md_dir = os.path.dirname(os.path.abspath(md_path))

    # ── Extract title + preamble ───────────────────────────────────────────────
    title         = ""
    kv: dict[str, str] = {}
    content_start = 0

    for idx, line in enumerate(lines):
        stripped = line.strip()
        if not title and stripped.startswith("# "):
            title = stripped[2:].strip()
            continue
        if title and re.match(r"^#{2,}", stripped):
            content_start = idx
            break
        if title:
            m = _KV_RE.match(stripped)
            if m:
                kv[m.group(1).strip()] = m.group(2).strip()
    else:
        content_start = len(lines)

    # ── Cover page ─────────────────────────────────────────────────────────────
    if title:
        _add_cover(doc, title, kv)

    # ── Body content ──────────────────────────────────────────────────────────
    i = content_start
    while i < len(lines):
        line     = lines[i]
        stripped = line.strip()

        # blank line
        if not stripped:
            i += 1
            continue

        # code fence
        if stripped.startswith("```"):
            i += 1
            code_lines: list[str] = []
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            i += 1
            _add_code_block(doc, code_lines)
            continue

        # heading (all levels handled; H4+ maps to Heading 3)
        m = re.match(r"^(#{1,6})\s+(.*)", stripped)
        if m:
            level = min(len(m.group(1)), 3)
            p     = doc.add_paragraph(style=f"Heading {level}")
            _add_runs(p, _strip_links(m.group(2)))
            i += 1
            continue

        # horizontal rule — skip
        if re.match(r"^[-*_]{3,}$", stripped):
            i += 1
            continue

        # image
        img_m = _IMAGE_RE.fullmatch(stripped)
        if img_m:
            img_path = os.path.join(md_dir, img_m.group(2))
            if os.path.exists(img_path):
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.add_run().add_picture(img_path, width=Inches(5.5))
            i += 1
            continue

        # table
        if stripped.startswith("|") and stripped.endswith("|"):
            tbl_lines: list[str] = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                tbl_lines.append(lines[i].strip())
                i += 1
            if len(tbl_lines) >= 2:
                headers  = [c.strip() for c in tbl_lines[0].strip("|").split("|")]
                from_row = 2 if _TABLE_SEP.match(tbl_lines[1]) else 1
                rows     = [
                    [c.strip() for c in row.strip("|").split("|")]
                    for row in tbl_lines[from_row:]
                ]
                _add_table(doc, headers, rows)
            continue

        # bullet list
        m = re.match(r"^(\s*)[-*+]\s+(.*)", line)
        if m:
            depth = len(m.group(1)) // 2
            p     = doc.add_paragraph(style="List Paragraph")
            p.paragraph_format.left_indent       = Inches(0.25 + 0.25 * depth)
            p.paragraph_format.first_line_indent = Inches(-0.18)
            p.paragraph_format.space_before      = Pt(1)
            p.paragraph_format.space_after       = Pt(1)
            p.add_run(("•" if depth == 0 else "○") + "  ")
            _add_runs(p, _strip_links(m.group(2)))
            i += 1
            continue

        # numbered list
        m = re.match(r"^\s*\d+\.\s+(.*)", line)
        if m:
            p = doc.add_paragraph(style="List Paragraph")
            p.paragraph_format.left_indent  = Inches(0.25)
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after  = Pt(1)
            _add_runs(p, _strip_links(m.group(1)))
            i += 1
            continue

        # blockquote
        if stripped.startswith(">"):
            p = doc.add_paragraph(style="Normal")
            p.paragraph_format.left_indent = Inches(0.4)
            r = p.add_run(stripped.lstrip(">").strip())
            r.italic = True; r.font.size = Pt(9)
            i += 1
            continue

        # regular paragraph
        p = doc.add_paragraph(style="Normal")
        _add_runs(p, _strip_links(stripped))
        i += 1

    out = output_path or os.path.splitext(md_path)[0] + ".docx"
    doc.save(out)
    print(f"✓  {os.path.relpath(md_path)} → {os.path.relpath(out)}")
    return out


# ── Named targets ──────────────────────────────────────────────────────────────

def _rq(rq: str, entregable: str, name: str) -> str:
    folder = os.path.join(_DOCS_DIR, rq, entregable)
    os.makedirs(folder, exist_ok=True)
    return os.path.join(folder, name)


NAMED: dict[str, str] = {
    "inventario_datasets": _rq("rq05", "E2", "Inventario_Fichas_Dataset_RQ05.md"),
    "matriz_linaje":       _rq("rq05", "E3", "Matriz_Linaje_Visual_RQ05.md"),
    "guia_stewards":       _rq("rq05", "E4", "Guia_Gobernanza_DataStewards_RQ05.md"),
}

DEFAULT_REPORTS: list[str] = list(NAMED.values())


# ── Entry point ────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    args = sys.argv[1:]

    if not args:
        paths = DEFAULT_REPORTS
    else:
        paths = []
        for a in args:
            if a in NAMED:
                paths.append(NAMED[a])
            else:
                paths.append(os.path.abspath(a))

    ok = err = 0
    for path in paths:
        if not os.path.exists(path):
            print(f"✗  not found: {path}"); err += 1; continue
        try:
            convert(path)
            ok += 1
        except Exception as exc:
            print(f"✗  {os.path.basename(path)}: {exc}"); err += 1

    print(f"\n{ok} converted, {err} failed.")

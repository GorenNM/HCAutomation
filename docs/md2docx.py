#!/usr/bin/env python3
"""Convert a Markdown file to .docx. No template, no project-specific content.

Usage:
    python docs/md2docx.py DOCUMENTACION.md            # -> DOCUMENTACION.docx
    python docs/md2docx.py in.md out.docx              # explicit output
    python docs/md2docx.py a.md b.md c.md              # several files

Handles: headings, tables, fenced code, bullet/numbered lists, blockquotes,
images, horizontal rules, and inline **bold** *italic* `code` ~~strike~~ [links].
"""

from __future__ import annotations

import os
import re
import sys

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

_BLACK = RGBColor(0x00, 0x00, 0x00)
MONO = "Consolas"  # ponytail: box-drawing chars render; swap to Courier New if unavailable

# ── Patterns ───────────────────────────────────────────────────────────────────
_INLINE = re.compile(
    r"\*\*(.+?)\*\*"
    r"|\*(.+?)\*"
    r"|`(.+?)`"
    r"|~~(.+?)~~"
)
_TABLE_SEP = re.compile(r"^\|[-|\s:]+\|$")
_IMAGE_RE  = re.compile(r"!\[([^\]]*)\]\(<?([^>)]+?)>?\)")
_LINK_RE   = re.compile(r"\[([^\]]+)\]\([^)]+\)")


# ── Inline formatting ──────────────────────────────────────────────────────────

def _add_runs(para, text: str) -> None:
    pos = 0
    for m in _INLINE.finditer(text):
        if m.start() > pos:
            para.add_run(text[pos:m.start()])
        bold, italic, code, strike = m.group(1), m.group(2), m.group(3), m.group(4)
        if bold:
            r = para.add_run(bold);   r.bold = True
        elif italic:
            r = para.add_run(italic); r.italic = True
        elif code:
            r = para.add_run(code);   r.font.name = MONO; r.font.size = Pt(9)
        elif strike:
            r = para.add_run(strike); r.font.strike = True
        pos = m.end()
    if pos < len(text):
        para.add_run(text[pos:])


def _strip_links(text: str) -> str:
    return _LINK_RE.sub(r"\1", text)


# ── Table helpers ──────────────────────────────────────────────────────────────

def _cell_bg(cell, hex_color: str) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def _cell_borders(cell) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    bdr  = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"),   "single")
        el.set(qn("w:sz"),    "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "000000")
        bdr.append(el)
    tcPr.append(bdr)


def _set_table_borders(tbl) -> None:
    tbl_el = tbl._tbl
    tblPr  = tbl_el.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl_el.insert(0, tblPr)
    tblBorders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"),   "single")
        el.set(qn("w:sz"),    "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "000000")
        tblBorders.append(el)
    tblPr.append(tblBorders)


def _repeat_header(row) -> None:
    """Mark a table row to repeat as header on every page."""
    row._tr.get_or_add_trPr().append(OxmlElement("w:tblHeader"))


def _add_table(doc, headers: list[str], rows: list[list[str]]) -> None:
    n = len(headers)
    tbl = doc.add_table(rows=1, cols=n)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    _set_table_borders(tbl)

    hdr_row = tbl.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        _cell_borders(cell)
        _cell_bg(cell, "DCDCDC")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(h.strip())
        r.bold = True; r.font.color.rgb = _BLACK; r.font.size = Pt(9)
    _repeat_header(hdr_row)

    for row in rows:
        tr = tbl.add_row()
        for i, val in enumerate(row[:n]):          # ponytail: extra cells dropped
            cell = tr.cells[i]
            _cell_borders(cell)
            p = cell.paragraphs[0]
            _add_runs(p, _strip_links(val.strip()))
            for r in p.runs:
                r.font.size = Pt(9)
        for i in range(len(row), n):               # short row -> pad borders
            _cell_borders(tr.cells[i])
    doc.add_paragraph()


def _add_code_block(doc, code_lines: list[str]) -> None:
    for line in code_lines:
        p = doc.add_paragraph(style="Normal")
        p.paragraph_format.left_indent  = Inches(0.3)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        run = p.add_run(line if line else " ")
        run.font.name = MONO; run.font.size = Pt(8.5)
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"),   "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"),  "F5F5F5")
        p._p.get_or_add_pPr().append(shd)
    doc.add_paragraph()


def _fix_styles(doc) -> None:
    """Black Calibri, sane heading sizes."""
    FONT = "Calibri"
    overrides = {
        "Title":          (24, True),
        "Heading 1":      (18, True),
        "Heading 2":      (13, True),
        "Heading 3":      (11, True),
        "Normal":         (10, False),
        "List Paragraph": (10, False),
    }
    for name, (size, bold) in overrides.items():
        try:
            s = doc.styles[name]
        except KeyError:
            continue
        s.font.name      = FONT
        s.font.size      = Pt(size)
        s.font.color.rgb = _BLACK
        s.font.bold      = bold or None


# ── Converter ──────────────────────────────────────────────────────────────────

def convert(md_path: str, output_path: str | None = None) -> str:
    with open(md_path, encoding="utf-8") as fh:
        lines = fh.read().splitlines()

    doc = Document()
    _fix_styles(doc)
    md_dir = os.path.dirname(os.path.abspath(md_path))

    i = 0
    while i < len(lines):
        line     = lines[i]
        stripped = line.strip()

        # blank line
        if not stripped:
            i += 1
            continue

        # code fence
        if stripped.startswith("```"):
            i += 1
            code_lines: list[str] = []
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            i += 1
            _add_code_block(doc, code_lines)
            continue

        # heading (H4+ maps to Heading 3)
        m = re.match(r"^(#{1,6})\s+(.*)", stripped)
        if m:
            level = min(len(m.group(1)), 3)
            p = doc.add_paragraph(style=f"Heading {level}")
            _add_runs(p, _strip_links(m.group(2)))
            i += 1
            continue

        # horizontal rule — skip
        if re.match(r"^[-*_]{3,}$", stripped):
            i += 1
            continue

        # image
        img_m = _IMAGE_RE.fullmatch(stripped)
        if img_m:
            img_path = os.path.join(md_dir, img_m.group(2))
            if os.path.exists(img_path):
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.add_run().add_picture(img_path, width=Inches(5.5))
            else:
                print(f"   ! image not found: {img_m.group(2)}")
            i += 1
            continue

        # table
        if stripped.startswith("|") and stripped.endswith("|"):
            tbl_lines: list[str] = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                tbl_lines.append(lines[i].strip())
                i += 1
            if len(tbl_lines) >= 2:
                headers  = [c.strip() for c in tbl_lines[0].strip("|").split("|")]
                from_row = 2 if _TABLE_SEP.match(tbl_lines[1]) else 1
                rows     = [
                    [c.strip() for c in row.strip("|").split("|")]
                    for row in tbl_lines[from_row:]
                ]
                _add_table(doc, headers, rows)
            continue

        # bullet list
        m = re.match(r"^(\s*)[-*+]\s+(.*)", line)
        if m:
            depth = len(m.group(1)) // 2
            p = doc.add_paragraph(style="List Paragraph")
            p.paragraph_format.left_indent       = Inches(0.25 + 0.25 * depth)
            p.paragraph_format.first_line_indent = Inches(-0.18)
            p.paragraph_format.space_before      = Pt(1)
            p.paragraph_format.space_after       = Pt(1)
            p.add_run(("•" if depth == 0 else "○") + "  ")
            _add_runs(p, _strip_links(m.group(2)))
            i += 1
            continue

        # numbered list — keep the literal number, markdown already has it
        m = re.match(r"^(\s*)(\d+\.)\s+(.*)", line)
        if m:
            depth = len(m.group(1)) // 2
            p = doc.add_paragraph(style="List Paragraph")
            p.paragraph_format.left_indent       = Inches(0.25 + 0.25 * depth)
            p.paragraph_format.first_line_indent = Inches(-0.25)
            p.paragraph_format.space_before      = Pt(1)
            p.paragraph_format.space_after       = Pt(1)
            p.add_run(m.group(2) + "  ")
            _add_runs(p, _strip_links(m.group(3)))
            i += 1
            continue

        # blockquote — merge consecutive lines into one paragraph
        if stripped.startswith(">"):
            quote: list[str] = []
            while i < len(lines) and lines[i].strip().startswith(">"):
                quote.append(lines[i].strip().lstrip(">").strip())
                i += 1
            p = doc.add_paragraph(style="Normal")
            p.paragraph_format.left_indent = Inches(0.4)
            _add_runs(p, _strip_links(" ".join(quote)))
            for r in p.runs:
                r.italic = True; r.font.size = Pt(9)
            continue

        # paragraph — merge until blank line / next block
        para: list[str] = []
        while i < len(lines):
            nxt = lines[i].strip()
            if not nxt or re.match(r"^(#{1,6}\s|```|\||>|[-*+]\s|\d+\.\s|[-*_]{3,}$)", nxt):
                break
            para.append(nxt)
            i += 1
        p = doc.add_paragraph(style="Normal")
        _add_runs(p, _strip_links(" ".join(para)))

    out = output_path or os.path.splitext(md_path)[0] + ".docx"
    doc.save(out)
    print(f"✓  {md_path} → {out}")
    return out


# ── Entry point ────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    args = sys.argv[1:]
    if not args:
        sys.exit(__doc__)

    if len(args) == 2 and args[1].lower().endswith(".docx"):
        jobs = [(args[0], args[1])]
    else:
        jobs = [(a, None) for a in args]

    ok = err = 0
    for src, dst in jobs:
        if not os.path.exists(src):
            print(f"✗  not found: {src}"); err += 1; continue
        try:
            convert(src, dst)
            ok += 1
        except Exception as exc:
            print(f"✗  {src}: {exc}"); err += 1

    print(f"\n{ok} converted, {err} failed.")
    sys.exit(1 if err else 0)
